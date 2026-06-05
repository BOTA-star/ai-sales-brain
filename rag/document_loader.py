from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

from pypdf import PdfReader
from docx import Document as DocxDocument


@dataclass
class LoadedDocument:
    content: str
    source: str
    file_name: str
    file_type: str
    page: Optional[int] = None

class DocumentLoader:
    def __init__(self, raw_data_dir: str = "data/raw"):
        self.raw_data_dir = Path(raw_data_dir)

    def load_all_documents(self) -> List[LoadedDocument]:
        documents: List[LoadedDocument] = []

        if not self.raw_data_dir.exists():
            raise FileNotFoundError(f"Folder not found: {self.raw_data_dir}")

        for file_path in self.raw_data_dir.iterdir():
            if file_path.is_file():
                documents.extend(self.load_document(file_path))

        return documents

    def load_document(self, file_path: Path) -> List[LoadedDocument]:
        suffix = file_path.suffix.lower()

        if suffix == ".txt":
            return self._load_txt(file_path)

        if suffix == ".md":
            return self._load_txt(file_path)

        if suffix == ".pdf":
            return self._load_pdf(file_path)

        if suffix == ".docx":
            return self._load_docx(file_path)

        print(f"Unsupported file type: {file_path.name}")
        return []

    def _load_txt(self, file_path: Path) -> List[LoadedDocument]:
        content = file_path.read_text(encoding="utf-8", errors="ignore")

        return [
            LoadedDocument(
                content=self._clean_text(content),
                source=str(file_path),
                file_name=file_path.name,
                file_type=file_path.suffix.lower(),
            )
        ]

    def _load_pdf(self, file_path: Path) -> List[LoadedDocument]:
        documents: List[LoadedDocument] = []

        reader = PdfReader(str(file_path))

        for page_index, page in enumerate(reader.pages):
            text = page.extract_text() or ""

            if text.strip():
                documents.append(
                    LoadedDocument(
                        content=self._clean_text(text),
                        source=str(file_path),
                        file_name=file_path.name,
                        file_type=".pdf",
                        page=page_index + 1,
                    )
                )

        return documents

    def _load_docx(self, file_path: Path) -> List[LoadedDocument]:
        doc = DocxDocument(str(file_path))

        paragraphs = [
            paragraph.text.strip()
            for paragraph in doc.paragraphs
            if paragraph.text.strip()
        ]

        content = "\n".join(paragraphs)

        return [
            LoadedDocument(
                content=self._clean_text(content),
                source=str(file_path),
                file_name=file_path.name,
                file_type=".docx",
            )
        ]

    def _clean_text(self, text: str) -> str:
        lines = [line.strip() for line in text.splitlines()]
        lines = [line for line in lines if line]
        return "\n".join(lines)